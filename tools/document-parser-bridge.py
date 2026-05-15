#!/usr/bin/env python3
"""Optional R3MES document parser bridge.

This script is intentionally dependency-light from the repo perspective:
install parser packages in an external/local venv, then point
R3MES_DOCUMENT_PARSER_COMMAND at that venv's python executable and use this
script as the command argument.

Supported when optional packages are installed:
- PDF: pypdf
- DOCX: python-docx
"""

from __future__ import annotations

import argparse
import html
import json
import re
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")


def fail(message: str, code: int = 2) -> None:
    print(message, file=sys.stderr)
    raise SystemExit(code)


def normalize_text(value: str) -> str:
    lines = [line.rstrip() for line in value.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return "\n".join(lines).strip()


def artifact(kind: str, text: str, *, title: str | None = None, page: int | None = None, level: int | None = None, items: list[str] | None = None, score: int | None = None) -> dict:
    clean = normalize_text(text)
    base_score = {
        "definition": 92,
        "table": 84,
        "list": 82,
        "qa": 80,
        "paragraph": 68,
        "heading": 30,
        "title": 24,
        "url": 1,
        "footer": 3,
        "page_marker": 2,
    }.get(kind, 45)
    return {
        "kind": kind,
        "text": clean,
        "title": title,
        "page": page,
        "level": level,
        "items": items,
        "answerabilityScore": score if score is not None else base_score,
    }


def classify_block(text: str) -> str:
    low = text.lower()
    stripped = text.strip()
    if stripped.startswith("http://") or stripped.startswith("https://"):
        return "url"
    if len(stripped) <= 90 and stripped.lower().startswith("hafta "):
        return "footer"
    if len(stripped) <= 140 and any(month in low for month in ["ocak", "şubat", "subat", "mart", "nisan", "mayıs", "mayis", "haziran", "temmuz", "ağustos", "agustos", "eylül", "eylul", "ekim", "kasım", "kasim", "aralık", "aralik"]) and any(ch.isdigit() for ch in stripped):
        return "footer"
    stripped_without_decimal_dots = re.sub(r"\d+\.\d+", "", stripped)
    if (
        len(stripped) <= 100
        and "\n" not in stripped
        and not any(mark in stripped_without_decimal_dots for mark in [".", "?", "!", ";", ":"])
    ):
        return "heading"
    if "|" in stripped and "\n" in stripped:
        return "table"
    if any(marker in low for marker in [" nedir", " denir", " ifade eder", " tanımlanır", " tanimlanir", " bütünüdür", " butunudur"]):
        return "definition"
    if stripped.count("\n- ") >= 2 or stripped.count("\n• ") >= 2 or stripped.count("\n") >= 2:
        return "list"
    return "paragraph"


def artifacts_from_markdown(markdown: str) -> list[dict]:
    artifacts: list[dict] = []
    current_title: str | None = None
    current_page: int | None = None
    block: list[str] = []

    def flush() -> None:
        nonlocal block
        text = normalize_text("\n".join(block))
        block = []
        if not text:
            return
        artifacts.append(artifact(classify_block(text), text, title=current_title, page=current_page))

    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            flush()
            continue
        if line.lower().startswith("## page "):
            flush()
            try:
                current_page = int(line.split()[-1])
            except Exception:
                current_page = None
            artifacts.append(artifact("page_marker", line.replace("## ", ""), title=current_title, page=current_page, score=2))
            continue
        if line.startswith("#"):
            flush()
            level = len(line) - len(line.lstrip("#"))
            title = line.lstrip("#").strip()
            current_title = title
            artifacts.append(artifact("title" if level == 1 else "heading", title, title=title, page=current_page, level=level))
            continue
        block.append(line)
    flush()
    return [item for item in artifacts if item.get("text")]


def parse_pdf(path: Path) -> tuple[str, list[dict]]:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:  # pragma: no cover - depends on optional env
        fail(
            "Missing optional dependency 'pypdf'. Install in an isolated venv, "
            "for example: python -m pip install pypdf",
        )

    reader = PdfReader(str(path))
    sections: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = normalize_text(page.extract_text() or "")
        if not text:
            continue
        sections.append(f"## Page {index}\n\n{text}")
    markdown = "\n\n".join(sections)
    return markdown, artifacts_from_markdown(markdown)


def markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [[html.escape(cell.strip()).replace("\n", " ") for cell in row] + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    body = normalized[1:]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def parse_docx(path: Path) -> tuple[str, list[dict]]:
    try:
        from docx import Document  # type: ignore
    except Exception:  # pragma: no cover - depends on optional env
        fail(
            "Missing optional dependency 'python-docx'. Install in an isolated venv, "
            "for example: python -m pip install python-docx",
        )

    doc = Document(str(path))
    sections: list[str] = []
    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            sections.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        table_md = markdown_table(rows)
        if table_md:
            sections.append(f"## Table {table_index}\n\n{table_md}")

    parsed = "\n\n".join(sections)
    xml_fallback = parse_docx_xml_text(path)
    if len(normalize_text(xml_fallback)) > len(normalize_text(parsed)) * 2:
        sections.append("## XML Text Fallback")
        sections.append(xml_fallback)
    markdown = "\n\n".join(sections)
    return markdown, artifacts_from_markdown(markdown)


def parse_docx_xml_text(path: Path) -> str:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    parts: list[str] = []
    try:
        with zipfile.ZipFile(path) as archive:
            names = [
                name
                for name in archive.namelist()
                if name.startswith("word/") and name.endswith(".xml") and not name.endswith(".rels")
            ]
            for name in sorted(names):
                try:
                    root = ET.fromstring(archive.read(name))
                except ET.ParseError:
                    continue
                texts = [
                    node.text.strip()
                    for node in root.findall(".//w:t", namespace)
                    if node.text and node.text.strip()
                ]
                if texts:
                    parts.append(f"### {name}\n\n" + "\n".join(texts))
    except zipfile.BadZipFile:
        return ""
    return normalize_text("\n\n".join(parts))


def main() -> None:
    parser = argparse.ArgumentParser(description="Parse PDF/DOCX into Markdown for R3MES ingestion.")
    parser.add_argument("input", help="Path to PDF or DOCX file")
    args = parser.parse_args()

    path = Path(args.input).resolve()
    if not path.exists():
        fail(f"Input file not found: {path}", 1)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        output, artifacts = parse_pdf(path)
        source_type = "PDF"
    elif suffix == ".docx":
        output, artifacts = parse_docx(path)
        source_type = "DOCX"
    else:
        fail(f"Unsupported bridge file type: {suffix}. Expected .pdf or .docx", 1)

    output = normalize_text(output)
    if not output:
        fail("Parser produced empty output", 3)
    print(json.dumps({
        "sourceType": source_type,
        "text": output,
        "artifacts": artifacts,
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
